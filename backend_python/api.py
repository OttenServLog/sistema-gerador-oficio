from flask import Flask, request, jsonify, send_file
import fitz  # PyMuPDF
import re
import os
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from babel.dates import format_date
from datetime import datetime
import tempfile

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
MODELO_PATH = 'modelo_oficio.docx'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def extract_data_from_pdf(filepath):
    with fitz.open(filepath) as doc:
        text = ""
        for page in doc:
            text += page.get_text()

    contas = re.findall(r'(\d{10} ?- ?\d)', text)
    conta_debito = contas[-1].replace(' ', '') if contas else 'Desconhecida'

    fornecedores = []
    fonte_alerta = False

    blocos = re.findall(r'EMPENHO\s+(.*?)(?=EMPENHO\s+|$)', text, re.DOTALL | re.IGNORECASE)

    for bloco in blocos:
        try:
            doc_id_match = re.search(r'\d{3}\.\d{3}\.\d{3}-\d{2}|\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}', bloco)
            if not doc_id_match:
                continue
            doc_id = doc_id_match.group(0)

            nome_bruto = bloco.split(doc_id)[0].strip().replace('\n', ' ')
            nome_limpo = re.sub(r'^\d+(\s+\d+)*\s+', '', nome_bruto)

            fonte_match = re.findall(r'\n(\d{1}\.\d{3}\.\d{3})\n', bloco)
            fonte = fonte_match[0] if fonte_match else ''
            if fonte.startswith("2"):
                fonte_alerta = True

            banco = re.search(r'Banco:\s*(\d+)', bloco)
            agencia = re.search(r'AG:\s*([\d\-]+)', bloco)
            conta = re.search(r'CC:\s*([\d\-]+)', bloco)

            parte_util = bloco.split("Banco:")[0]
            valores = re.findall(r'(\d{1,3}(?:\.\d{3})*,\d{2})', parte_util)

            if len(valores) < 2:
                continue

            valor_desc = valores[-2]
            valor_liq = valores[-1]

            fornecedores.append({
                'nome': nome_limpo,
                'cnpj': doc_id,
                'banco': banco.group(1) if banco else '',
                'agencia': agencia.group(1) if agencia else '',
                'conta': conta.group(1),
                'valorLiquido': valor_liq,
                'fonte': fonte,
                'desconto': valor_desc
            })

            if valor_desc != '0,00':
                fornecedores.append({
                    'nome': 'PREFEITURA MUNICIPAL DE UBERABA',
                    'cnpj': '18.428.839/0001-90',
                    'banco': '001',
                    'agencia': '0015-9',
                    'conta': '118.252-8',
                    'valorLiquido': valor_desc,
                    'fonte': '',
                    'desconto': ''
                })

        except Exception as e:
            print(f"Erro ao processar fornecedor: {e}")
            continue

    return conta_debito, fornecedores, fonte_alerta, text

@app.route('/upload', methods=['POST'])
def upload_pdf():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({'erro': 'Arquivo inválido.'}), 400

    filepath = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(filepath)

    conta, fornecedores, fonte_alerta, texto_pdf = extract_data_from_pdf(filepath)

    return jsonify({
        'conta_debito': conta,
        'fornecedores': fornecedores,
        'fonte_alerta': fonte_alerta,
        'texto_pdf': texto_pdf
    })

@app.route('/gerar-oficio', methods=['POST'])
def gerar_oficio():
    data = request.json
    numero = data.get('numeroOficio')
    assinatura1 = data.get('assinatura1')
    assinatura2 = data.get('assinatura2')
    tabelas = data.get('tabelas')
    assinaturas = data.get('assinaturas')

    doc = Document(MODELO_PATH)

    data_formatada = format_date(datetime.now(), format="dd 'de' MMMM 'de' yyyy", locale='pt_BR')
    for p in doc.paragraphs:
        if "____/____" in p.text:
            p.text = p.text.replace("____/____", numero)
        if "___/___/____" in p.text:
            p.text = p.text.replace("___/___/____", data_formatada)

    doc.add_paragraph()

    for idx, grupo in enumerate(tabelas):
        conta = grupo.get("conta")
        fornecedores = grupo.get("fornecedores", [])

        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'

        widths = [Inches(2.5), Inches(1.7), Inches(2.3), Inches(1.3)]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = widths[i]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        cell = table.rows[0].cells
        cell[0].merge(cell[3])
        cell[0].text = f'CONTA DO DÉBITO: {conta}'
        cell[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        headers = ['FAVORECIDO', 'CNPJ/CPF', 'CONTA PARA CRÉDITO', 'VALOR']
        for i, h in enumerate(headers):
            cell = table.rows[1].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

        for f in fornecedores:
            row = table.add_row().cells
            row[0].text = f['nome']
            row[1].text = f['cnpj']
            row[2].text = f"Banco: {f['banco']}\nAgência: {f['agencia']}\nConta: {f['conta']}"
            row[3].text = f"R$ {f['valorLiquido']}"
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)

        if idx < len(tabelas) - 1:
            doc.add_paragraph()  # espaçamento entre blocos de conta

    doc.add_paragraph()

    def adicionar_assinatura(nome_assinatura):
        dados = next((a for a in assinaturas if a['nome'] == nome_assinatura), None)
        if not dados:
            return
        par = doc.add_paragraph()
        run_nome = par.add_run(dados['nome'] + "\n")
        run_nome.bold = True
        par.add_run(dados['cargo'] + "\n")
        par.add_run(dados['decreto'])
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    adicionar_assinatura(assinatura1)
    adicionar_assinatura(assinatura2)

    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        return send_file(tmp.name, as_attachment=True, download_name=f'OFICIO {numero} - AUTORIZAÇÃO DE PAGAMENTO.docx')

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
